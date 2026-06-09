"""
TalentMind AI — PDF / DOCX Text Extractor
Supports .pdf (via pdfminer.six) and .docx/.doc (via python-docx).
"""
import io
import re
from pathlib import Path

from pdfminer.high_level import extract_text_to_fp
from pdfminer.layout import LAParams
from docx import Document
import aiofiles


async def extract_text_from_file(file_path: str) -> str:
    """
    Asynchronously extract raw text from a PDF or DOCX file.

    Parameters
    ----------
    file_path:
        Absolute path to the uploaded file.

    Returns
    -------
    str
        Raw extracted text, cleaned of excessive whitespace.

    Raises
    ------
    ValueError
        For unsupported file extensions.
    """
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        text = extract_pdf_text(file_path)
    elif suffix in (".docx", ".doc"):
        text = extract_docx_text(file_path)
    else:
        raise ValueError(f"Unsupported file type: {suffix!r}")

    return clean_text(text)


def extract_pdf_text(file_path: str) -> str:
    """
    Extract text from a PDF using pdfminer's high-level API.
    Handles multi-column layouts via LAParams heuristics.
    """
    output = io.StringIO()
    laparams = LAParams(
        line_margin=0.5,
        word_margin=0.1,
        char_margin=2.0,
        boxes_flow=0.5,
        detect_vertical=False,
    )
    with open(file_path, "rb") as f:
        extract_text_to_fp(f, output, laparams=laparams, output_type="text", codec="utf-8")
    return output.getvalue()


def extract_docx_text(file_path: str) -> str:
    """
    Extract all text from a DOCX file including table cells and headers/footers.
    """
    doc = Document(file_path)
    parts: list[str] = []

    # Main paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Table cells
    for table in doc.tables:
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_cells:
                parts.append("  ".join(row_cells))

    # Headers and footers
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            if paragraph.text.strip():
                parts.append(paragraph.text.strip())
        for paragraph in section.footer.paragraphs:
            if paragraph.text.strip():
                parts.append(paragraph.text.strip())

    return "\n".join(parts)


def clean_text(text: str) -> str:
    """
    Normalise whitespace and strip non-ASCII characters from extracted text.
    """
    # Collapse multiple blank lines to a single newline
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Normalise whitespace within lines
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.splitlines()]
    # Remove lines that are purely punctuation/noise
    lines = [ln for ln in lines if re.search(r"[a-zA-Z0-9]", ln)]
    return "\n".join(lines).strip()
